from io import BytesIO
from django.contrib.auth.decorators import user_passes_test
from django.shortcuts import render, redirect
from django.urls import reverse_lazy
from django.views.generic import ListView, DetailView, CreateView, TemplateView
from .models import RealEstate
import folium
from geopy.geocoders import Nominatim
from .filters import *
from .forms import DateTimeForm, AddingObjectForm, CreateContractForm
import docx
from django.core.mail import EmailMessage


# Create your views here.
class MainPage(ListView):
    context_object_name = 'real_estate_list'
    template_name = 'RealEstateApp/MainPage.html'

    def get_queryset(self):
        queryset = RealEstate.objects.all()
        status = self.request.GET.get('status')
        if status:
            queryset = queryset.filter(status=status)
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['rental_list'] = self.get_queryset().filter(category__name='Аренда').order_by('-pk')[:5]
        context['sale_list'] = self.get_queryset().filter(category__name='Продажа').order_by('-pk')[:5]
        return context


class DetailRealEstatePage(DetailView):
    model = RealEstate
    template_name = 'RealEstateApp/DetailRealEstate.html'
    context_object_name = 'object'

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = RealEstate.objects.get(pk=self.kwargs['pk']).address

        # map
        # Создание карты и изначального её расположения карты #
        m = folium.Map(location=[55.7887400, 49.1221400], zoom_start=12, min_zoom=6)
        popup = folium.Popup(
            'Город: ' + str(self.object.city) + '</br></br>' + 'Район: ' + str(self.object.district) +
            '</br></br>' + 'Площадь: ' + str(self.object.area) + '</br></br>' + 'Этаж: ' + str(self.object.floor) +
            '</br></br>' + 'Кол-во комнат: ' + str(
                self.object.rooms) + '</br></br>' + 'Планировка: ' + self.object.layout +
            '</br></br>' + 'Ремонт: ' + self.object.finishing + '</br></br>' + 'Сан.узел: ' + self.object.bathroom +
            '</br></br>' + 'Балкон: ' + self.object.balcony + '</br></br>', max_width=300)

        # getting geolocation
        geolocator = Nominatim(user_agent="RealEstate")
        folium.Marker(location=[self.object.latitude, self.object.longitude], tooltip='Нажмите чтобы узнать больше',
                      popup=popup).add_to(m)
        m = m._repr_html_()
        context['map'] = m
        return context


class RentalRealEstatesPage(ListView):
    model = RealEstate
    context_object_name = 'rental_objects'
    template_name = 'RealEstateApp/RentalListPage.html'
    paginate_by = 4

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['rental_filter'] = RentalFilter(self.request.GET,
                                                queryset=RealEstate.objects.filter(category__name='Аренда').order_by(
                                                    '-pk'))

        m = folium.Map(location=[55.7887400, 49.1221400], zoom_start=12, min_zoom=6)
        for object in RentalFilter(self.request.GET,
                                   queryset=RealEstate.objects.filter(category__name='Аренда').order_by('-pk')).qs:
            popup = folium.Popup(
                'Город: ' + str(object.city) + '</br></br>' + 'Район: ' + str(object.district) +
                '</br></br>' + 'Площадь: ' + str(object.area) + '</br></br>' + 'Этаж: ' + str(object.floor) +
                '</br></br>' + 'Кол-во комнат: ' + str(
                    object.rooms) + '</br></br>' + 'Планировка: ' + object.layout +
                '</br></br>' + 'Ремонт: ' + object.finishing + '</br></br>' + 'Сан.узел: ' + object.bathroom +
                '</br></br>' + 'Балкон: ' + object.balcony + '</br></br>', max_width=300)

            geolocator = Nominatim(user_agent="RealEstate")

            folium.Marker(location=[object.latitude, object.longitude], tooltip='Нажмите чтобы узнать больше',
                          popup=popup).add_to(m)

        m = m._repr_html_()
        context['map'] = m
        return context

    def get_queryset(self):
        queryset = super().get_queryset().filter(category__name='Аренда').order_by('-pk')
        return RentalFilter(self.request.GET, queryset=queryset).qs


class SellingRealEstatesPage(ListView):
    model = RealEstate
    template_name = 'RealEstateApp/SellingListPage.html'
    paginate_by = 10
    context_object_name = 'selling_objects'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['selling_filter'] = RentalFilter(self.request.GET,
                                                 queryset=RealEstate.objects.filter(category__name='Продажа').order_by(
                                                     '-pk'))
        m = folium.Map(location=[55.7887400, 49.1221400], zoom_start=12, min_zoom=6)
        for object in RentalFilter(self.request.GET,
                                   queryset=RealEstate.objects.filter(category__name='Продажа').order_by('-pk')).qs:
            popup = folium.Popup(
                'Город: ' + str(object.city) + '</br></br>' + 'Район: ' + str(object.district) +
                '</br></br>' + 'Площадь: ' + str(object.area) + '</br></br>' + 'Этаж: ' + str(object.floor) +
                '</br></br>' + 'Кол-во комнат: ' + str(
                    object.rooms) + '</br></br>' + 'Планировка: ' + object.layout +
                '</br></br>' + 'Ремонт: ' + object.finishing + '</br></br>' + 'Сан.узел: ' + object.bathroom +
                '</br></br>' + 'Балкон: ' + object.balcony + '</br></br>', max_width=300)

            geolocator = Nominatim(user_agent="RealEstate")
            folium.Marker(location=[object.latitude, object.longitude], tooltip='Нажмите чтобы узнать больше',
                          popup=popup).add_to(m)

        m = m._repr_html_()
        context['map'] = m

        return context

    def get_queryset(self):
        queryset = super().get_queryset().filter(category__name='Продажа').order_by('-pk')
        return RentalFilter(self.request.GET, queryset=queryset).qs


class BookingObject(DetailView):
    model = RealEstate
    template_name = 'RealEstateApp/BookingObject.html'
    context_object_name = 'object'
    form_class = DateTimeForm  # Добавляем кастомную форму

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['form'] = DateTimeForm()  # Добавляем форму в контекст
        return context

    def post(self, request, pk):
        if request.method == 'POST':
            form = DateTimeForm(request.POST)
            if form.is_valid():
                time_slot = TimeSlot.objects.create(time=form.cleaned_data['time'])
                booking = Booking.objects.create(object=self.get_object(), time_slot=time_slot,
                                                 client_name=request.user, realtor=self.get_object().realtor)
                return redirect('confirmation')
            else:
                form = DateTimeForm()
                return redirect('main_page')


class BookingConfirmation(TemplateView):
    template_name = 'RealEstateApp/BookingConfirmation.html'


class ContactsPage(TemplateView):
    template_name = 'RealEstateApp/ContactsPage.html'


class AddingObject(CreateView):
    model = RealEstate
    form_class = AddingObjectForm
    template_name = 'RealEstateApp/AddingObject.html'
    success_url = reverse_lazy('main_page')


class CreateContract(CreateView):
    model = Contracts
    form_class = CreateContractForm
    template_name = 'RealEstateApp/CreateContract.html'
    success_url = reverse_lazy('main_page')

    def form_valid(self, form):
        # Автоматическое заполнение поля стоимости в форме перед сохранением
        form.instance.price = float(form.instance.object.price) * 1.2

        # Создание Word-документа
        doc = docx.Document()
        # Получение данных из базы
        contract = form.save(commit=False)  # Сохраняем форму, но не сохраняем в базу данных пока
        object = contract.object
        client_name = contract.client_name
        realtor = contract.realtor
        price = float(object.price) * 1.2
        date = contract.date
        print(object, client_name, realtor, price, date)
        print('Заполнение Word-документа данными')
        # Заполнение Word-документа данными
        doc.add_paragraph(f"Объект недвижимости: {object}")
        doc.add_paragraph(f"Клиент: {client_name}")
        doc.add_paragraph(f"Риелтор: {realtor}")
        doc.add_paragraph(f"Стоимость: {price}")
        doc.add_paragraph(f"Дата: {date}")
        print('Сохранение Word-документа в памяти')
        # Сохранение Word-документа в памяти
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        print('Отправка Word-документа на почту')
        # Отправка Word-документа на почту
        email = EmailMessage(
            subject='Договор',
            body='Пожалуйста, найдите прикрепленный договор.',
            from_email='dmitry.solonko.2001@yandex.ru',
            to=[client_name.email],  # Здесь использовано поле email из модели клиент
        )
        email.attach('contract.docx', doc_buffer.getvalue(),
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        email.send()
        contract.save()  # Теперь сохраняем объект контракта в базу данных
        return super().form_valid(form)

